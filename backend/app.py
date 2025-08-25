# app.py
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from dotenv import load_dotenv
import os, json, hashlib

# ---------------------------
# Load environment variables
# ---------------------------
load_dotenv()

# ---------------------------
# Try to import your graph + llm helper
# ---------------------------
# - run_graph: LangGraph pipeline compiled in validator.py
# - llm_complete: unified Gemini wrapper from agents/base.py (or base.py), with fake fallback
try:
    from validator import run_validation as run_graph  # your LangGraph compile wrapper
except Exception:
    run_graph = None

llm_complete = None
try:
    from agents.base import llm_complete as _llm
    llm_complete = _llm
except Exception:
    try:
        from base import llm_complete as _llm  # fallback if base.py isn't inside "agents/"
        llm_complete = _llm
    except Exception:
        pass

def _llm(text: str) -> str:
    """Safe LLM call with graceful fallback."""
    if callable(llm_complete):
        return llm_complete(text)
    return f"[FAKE GEMINI RESPONSE] {text[:200]}..."

# ---------------------------
# Initialize FastAPI App
# ---------------------------
app = FastAPI(title="Startup Validator API")

# ---------------------------
# CORS Configuration (dev-friendly)
# ---------------------------
allow_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:8501").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[o.strip() for o in allow_origins if o.strip()],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------
# Global Storage for Report
# ---------------------------
last_report = {}  # {"idea": str, "mode": "fast"|"deep", "report": {...}}

# ---------------------------
# Deterministic dynamic metrics (fallback when JSON parse fails or no key)
# ---------------------------
def _fallback_metrics(idea: str):
    # Per-idea deterministic numbers so different ideas ≠ same graph
    h = int(hashlib.sha256(idea.encode("utf-8")).hexdigest(), 16)
    base = 5 + (h % 10)                      # 5–14 ($k)
    growth = 0.12 + ((h >> 8) % 9) / 100     # 12–20%
    months = 12
    mrr = []
    val = base
    for m in range(months):
        wobble = 1 + (((h >> (m + 13)) & 3) - 1.5) * 0.02
        val = val * (1 + growth) * wobble
        mrr.append(round(val, 1))

    tam = 20 + (h % 80)                      # $20–$99B
    sam = round(tam * (0.3 + ((h >> 5) % 30) / 100), 1)  # 30–59% of TAM
    som = round(sam * (0.2 + ((h >> 11) % 20) / 100), 1) # 20–39% of SAM

    return {
        "problem": "Latency and manual workflows reduce conversion and cause revenue leakage.",
        "solution": "Automate the workflow with specialized agents; integrate with existing systems.",
        "trends": ["Agentic workflows", "LLM ops", "Vertical AI platforms", "Data governance"],
        "risks": ["Data quality", "Security & compliance", "Vendor lock-in", "Unit economics"],
        "market": {"TAM": tam, "SAM": sam, "SOM": som},
        "traction": {"monthly_mrr": mrr},
    }

def _structured_summary_with_llm(idea: str, market: str, competitors: str, financials: str):
    prompt = f"""
You are a startup analyst. Using the EVIDENCE below, return ONLY valid JSON (no prose).

EVIDENCE:
--- MARKET ---
{market}

--- COMPETITORS ---
{competitors}

--- FINANCIALS ---
{financials}

SCHEMA:
{{
  "problem": string,
  "solution": string,
  "trends": string[],
  "risks": string[],
  "market": {{ "TAM": number, "SAM": number, "SOM": number }},  // billions USD
  "traction": {{ "monthly_mrr": number[] }}                      // 6–12 values, thousands USD
}}

Rules:
- Provide 6–12 monthly_mrr values, roughly increasing (small plateaus OK).
- Market numbers are billions USD (floats allowed).
- Output ONLY JSON (no backticks, no explanations).
"""
    raw = _llm(prompt)
    try:
        data = json.loads(raw)
        # minimal sanity checks
        if not isinstance(data.get("traction", {}).get("monthly_mrr", []), list):
            raise ValueError("monthly_mrr missing")
        return data
    except Exception:
        return _fallback_metrics(idea)

def _sections_to_json(idea: str, market: str, competitors: str, financials: str):
    """
    Turn long agent paragraphs into structured JSON (headings + bullets).
    Falls back to a minimal JSON if the LLM/key isn't available.
    """
    # Minimal fallback
    if not (market or competitors or financials):
        return {
            "market": {"sections": []},
            "competitors": {"sections": []},
            "financials": {"sections": []},
        }

    prompt = f"""
You are a precise information extractor.
Convert the AGENT TEXT into STRICT JSON that captures headings and bullet points.

RETURN ONLY JSON (no prose, no backticks) matching this schema:

{{
  "market": {{
    "sections": [{{ "title": string, "bullets": string[] }}]
  }},
  "competitors": {{
    "sections": [{{ "title": string, "bullets": string[] }}]
  }},
  "financials": {{
    "sections": [{{ "title": string, "bullets": string[] }}]
  }}
}}

AGENT TEXT:
[MARKET]
{market}

[COMPETITORS]
{competitors}

[FINANCIALS]
{financials}
"""
    raw = _llm(prompt)
    try:
        data = json.loads(raw)
        # sanity shape
        for key in ["market", "competitors", "financials"]:
            if key not in data or "sections" not in data[key]:
                raise ValueError("bad sections")
        return data
    except Exception:
        # safe fallback shape that still satisfies JSON contract
        def to_sections(txt: str):
            if not txt:
                return []
            # fallback: single section with lines split
            lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
            return [{"title": "Summary", "bullets": lines[:20]}]

        return {
            "market": {"sections": to_sections(market)},
            "competitors": {"sections": to_sections(competitors)},
            "financials": {"sections": to_sections(financials)},
        }

# ---------------------------
# Unified validation function
# ---------------------------
def run_validation(idea: str, mode: str = "fast"):
    """
    1) Runs your LangGraph pipeline if available (market -> competitors -> financials -> report)
    2) Asks LLM for a structured JSON summary (problem/solution/trends/risks/market/traction)
    3) For deep mode, also returns `deep_json` (structured agent details)
    """
    # 1) Run LangGraph pipeline if present
    market = competitors = financials = report_text = ""
    if callable(run_graph):
        try:
            state = run_graph(idea, mode)  # expects keys: market, competitors, financials, report
            market = state.get("market", "") or ""
            competitors = state.get("competitors", "") or ""
            financials = state.get("financials", "") or ""
            report_text = state.get("report", "") or ""
        except Exception as e:
            report_text = f"(Graph error: {e})"

    # 2) Convert to structured fields
    summary = _structured_summary_with_llm(idea, market, competitors, financials)

    # 3) Deep JSON (only for deep mode)
    deep_json = None
    if str(mode).lower() == "deep":
        deep_json = _sections_to_json(idea, market, competitors, financials)

    # 4) Return shape expected by frontend (+ deep_json)
    return {
        "problem": summary.get("problem", ""),
        "solution": summary.get("solution", ""),
        "trends": summary.get("trends", []),
        "risks": summary.get("risks", []),
        "market": summary.get("market", {}),
        "traction": summary.get("traction", {}),
        "report_text": report_text,
        # long text (still available if you want to render it)
        "sections": {
            "market": market,
            "competitors": competitors,
            "financials": financials,
        },
        # new: structured JSON for deep view
        "deep_json": deep_json,
    }

# ---------------------------
# API Endpoints
# ---------------------------
@app.post("/validate")
async def validate(payload: dict):
    """Validate startup idea using fast or deep analysis."""
    global last_report
    idea = (payload.get("idea") or "").strip()
    mode = (payload.get("mode") or "fast").strip().lower()

    if not idea:
        return {"error": "Please enter a startup idea."}
    if mode not in {"fast", "deep"}:
        mode = "fast"

    final_report = run_validation(idea, mode)
    last_report = {"idea": idea, "report": final_report, "mode": mode}
    return last_report

@app.get("/generate_report")
def generate_report():
    """
    Generate a Word pitch deck.
    - If last_report exists, we use its content to shape slides.
    - Otherwise, ask LLM to draft a generic deck.
    Works even without an API key (falls back to fake content).
    """
    idea = last_report.get("idea", "Your Startup")
    report = last_report.get("report", {})
    problem = report.get("problem", "")
    solution = report.get("solution", "")
    trends = report.get("trends", [])
    risks = report.get("risks", [])
    market = report.get("market", {})
    traction = report.get("traction", {})
    narrative = report.get("report_text", "")

    # Build a prompt seeded with current findings
    pitch_deck_prompt = f"""
You are an expert startup strategist and pitch deck writer.
Create concise, investor-ready slide content with numbered bullets for each slide below.

Idea: {idea}

Current Findings:
- Problem: {problem}
- Solution: {solution}
- Trends: {", ".join(trends) if trends else "n/a"}
- Risks: {", ".join(risks) if risks else "n/a"}
- Market (billions): TAM={market.get('TAM','?')}, SAM={market.get('SAM','?')}, SOM={market.get('SOM','?')}
- Traction (k MRR): {traction.get('monthly_mrr', [])}
- Narrative: {narrative[:800]}

Slides:
1. Hook
2. Problem
3. Solution
4. Market
5. Business Model
6. Traction
7. Competition
8. Trends
9. Risks
10. Team
11. The Ask
12. Vision

Output format:
- Separate slides with a blank line.
- Each slide starts with the slide title on the first line.
- Then 3–6 short numbered bullets.
"""
    deck_content = _llm(pitch_deck_prompt).strip()

    # Create Word document
    doc = Document()
    doc.add_heading(f"Startup Pitch Deck: {idea}", 0)

    for section in deck_content.split("\n\n"):
        lines = [ln for ln in section.strip().split("\n") if ln.strip()]
        if not lines:
            continue
        title = lines[0].strip()
        if title:
            doc.add_heading(title, level=1)
        for line in lines[1:]:
            doc.add_paragraph(line.strip(), style="List Number")

    file_path = "Startup_Pitch_Deck.docx"
    doc.save(file_path)

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Startup_Pitch_Deck.docx",
    )
